import re
from dataclasses import dataclass
from pathlib import Path
from typing import List
from ingestion.cleaner import clean_text

def _doc_to_docx(doc_path: Path) -> Path:
    """
    Converts a legacy .doc file to .docx using Microsoft Word via COM (pywin32).
    Returns the path to the converted .docx file.
    Requires Microsoft Word to be installed on Windows.
    """
    import win32com.client
    docx_path = doc_path.with_suffix(".docx")
    if docx_path.exists():
        return docx_path
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        d = word.Documents.Open(str(doc_path.absolute()))
        d.SaveAs2(str(docx_path.absolute()), FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
        d.Close()
    finally:
        word.Quit()
    return docx_path

@dataclass
class ParsedPage:
    page_number: int
    text: str

def parse_docx(docx_path: Path) -> List[ParsedPage]:
    """
    Parses 3GPP .docx files using python-docx, preserving clause hierarchy and markdown tables.
    """
    import docx
    doc = docx.Document(docx_path)
    lines: List[str] = []

    for elem in doc.paragraphs:
        txt = elem.text.strip()
        if not txt:
            continue
        
        # Detect headings — 3GPP docx files can have corrupt style refs, guard against it
        try:
            style_name = elem.style.name.lower() if elem.style and getattr(elem.style, "name", None) else ""
        except Exception:
            style_name = ""
        if "heading 1" in style_name or "title" in style_name:
            lines.append(f"\n# {txt}\n")
        elif "heading 2" in style_name:
            lines.append(f"\n## {txt}\n")
        elif "heading 3" in style_name:
            lines.append(f"\n### {txt}\n")
        elif "heading 4" in style_name or "heading 5" in style_name:
            lines.append(f"\n#### {txt}\n")
        elif re.match(r'^\d+(\.\d+)+\s+', txt):
            # Numeric clause heading
            depth = min(4, txt.split()[0].count("."))
            prefix = "#" * (depth + 1)
            lines.append(f"\n{prefix} {txt}\n")
        else:
            lines.append(txt)

    for table in doc.tables:
        t_lines = []
        for row_idx, row in enumerate(table.rows):
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            t_lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                t_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        if t_lines:
            lines.append("\n" + "\n".join(t_lines) + "\n")

    full_text = "\n\n".join(lines)
    # Break into logical pseudo-pages of ~3000 chars
    chunk_size = 3000
    pages = []
    for i in range(0, max(1, len(full_text)), chunk_size):
        page_slice = full_text[i:i+chunk_size]
        pages.append(ParsedPage(page_number=(i // chunk_size) + 1, text=clean_text(page_slice)))

    return pages

def parse_pdf(pdf_path: Path) -> List[ParsedPage]:
    """
    Parses PDF using PyMuPDF and extracts formatted text per page.
    """
    import fitz
    doc = fitz.open(pdf_path)
    pages: List[ParsedPage] = []

    for page_idx in range(len(doc)):
        page = doc[page_idx]
        raw_text = page.get_text("text")
        cleaned = clean_text(raw_text)
        pages.append(ParsedPage(page_number=page_idx + 1, text=cleaned))

    doc.close()
    return pages

def parse_document(doc_path: Path) -> List[ParsedPage]:
    """
    Dispatcher to parse either .docx or .pdf.
    """
    ext = doc_path.suffix.lower()
    if ext == ".doc":
        doc_path = _doc_to_docx(doc_path)  # convert to .docx via MS Word COM
    if ext in [".docx", ".doc"]:
        return parse_docx(doc_path)
    elif ext == ".pdf":
        return parse_pdf(doc_path)
    else:
        # Fallback text read
        content = doc_path.read_text(encoding="utf-8", errors="ignore")
        return [ParsedPage(page_number=1, text=clean_text(content))]
